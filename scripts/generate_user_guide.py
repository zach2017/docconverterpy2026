#!/usr/bin/env python3
"""
Generate a comprehensive User Guide DOCX for the Document Conversion Service.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page setup ───────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ───────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, (size, color) in enumerate([
    (Pt(28), RGBColor(0x1a, 0x56, 0xc4)),  # Heading 1
    (Pt(20), RGBColor(0x24, 0x65, 0xb8)),  # Heading 2
    (Pt(14), RGBColor(0x2d, 0x73, 0xac)),  # Heading 3
], start=1):
    hs = doc.styles[f'Heading {level}']
    hs.font.size = size
    hs.font.color.rgb = color
    hs.font.bold = True
    hs.font.name = 'Calibri'
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    hs.paragraph_format.space_after = Pt(8)

def add_code_block(doc, text):
    """Add a monospace code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    # Shading
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F2F5')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1e, 0x1e, 0x2e)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Light Grid Accent 1'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    doc.add_paragraph()  # spacer
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xc0, 0x6a, 0x00)
    p.add_run(text)
    return p

# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Document Conversion Service")
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xc4)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Complete User Guide")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run("Version 1.0  •  2025")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc.paragraph_format.space_before = Pt(30)
run = desc.add_run(
    "A production-grade microservice that converts documents and images to plain text.\n"
    "Supports SQS, RabbitMQ, Kafka message buses, REST API, and 10 document formats."
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════

doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Overview",
    "2. Architecture",
    "3. Prerequisites",
    "4. Quick Start – Docker Compose",
    "5. Configuration Reference",
    "6. Message Schema",
    "7. REST API Reference",
    "8. Using the API Test Console",
    "9. Message Bus Integration",
    "    9.1 Amazon SQS",
    "    9.2 RabbitMQ",
    "    9.3 Apache Kafka",
    "10. Document Converters",
    "11. Image OCR",
    "12. File Fetchers (S3, URL, FTP)",
    "13. Output Storage (S3)",
    "14. Memory Efficiency & Streaming",
    "15. Testing",
    "16. Troubleshooting",
    "17. curl Examples Cookbook",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════

doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "The Document Conversion Service is a containerized Python microservice that accepts "
    "conversion jobs from multiple sources – three message buses (Amazon SQS, RabbitMQ, "
    "Apache Kafka) and an optional REST API – and converts documents and images into plain "
    "text.  The extracted text is uploaded to an Amazon S3 bucket (or S3-compatible storage "
    "like LocalStack or MinIO)."
)
doc.add_paragraph(
    "Every component is independently toggleable via environment variables.  You can run "
    "the service as a pure queue consumer with no HTTP surface, or as an API-only service, "
    "or any combination."
)

doc.add_heading("Key Features", level=3)
add_bullet(doc, "10 document formats: PDF, DOCX, XLSX, PPTX, HTML, RTF, ODT, TXT, CSV, Image (OCR)")
add_bullet(doc, "3 message buses: SQS, RabbitMQ, Kafka – all using the same JSON message schema")
add_bullet(doc, "REST API with file upload and JSON job submission")
add_bullet(doc, "Web-based API Test Console (built-in HTML UI)")
add_bullet(doc, "3 document sources: S3, HTTP/HTTPS URLs, FTP servers")
add_bullet(doc, "Authentication support: None, Basic, Bearer token")
add_bullet(doc, "Memory-efficient streaming via Python generators and temporary files")
add_bullet(doc, "PDF image extraction with OCR (Tesseract)")
add_bullet(doc, "Full Docker Compose stack with LocalStack, RabbitMQ, Kafka, FTP")
add_bullet(doc, "Every component can be enabled/disabled via environment variables")

# ═══════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("2. Architecture", level=1)

doc.add_heading("System Diagram", level=3)
doc.add_paragraph(
    "The service follows a pipeline pattern: Ingest → Fetch → Convert → Upload."
)
add_code_block(doc,
    "┌──────────────┐  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐\n"
    "│   SQS Queue  │  │  RabbitMQ Q  │  │  Kafka Topic │  │   REST API   │\n"
    "│ (LocalStack) │  │              │  │              │  │  (FastAPI)   │\n"
    "└──────┬───────┘  └──────┬───────┘  └──────┬───────┘  └──────┬───────┘\n"
    "       │                 │                 │                 │\n"
    "       └────────────────┬┘─────────────────┘                 │\n"
    "                        ▼                                    ▼\n"
    "              ┌─────────────────────────────────────────────────┐\n"
    "              │              Conversion Job Router               │\n"
    "              └──────────────────┬──────────────────────────────┘\n"
    "                                 ▼\n"
    "              ┌──────────────────────────────────────┐\n"
    "              │          Document Fetcher             │\n"
    "              │  S3  │  HTTP/URL  │  FTP  │  Local   │\n"
    "              └──────────────────┬───────────────────┘\n"
    "                                 ▼  (tmp file on disk)\n"
    "              ┌──────────────────────────────────────┐\n"
    "              │        Converter Dispatcher           │\n"
    "              │  PDF │ DOCX │ XLSX │ PPTX │ HTML ... │\n"
    "              └──────────────────┬───────────────────┘\n"
    "                                 ▼  (text chunks via yield)\n"
    "              ┌──────────────────────────────────────┐\n"
    "              │      S3 Output Storage                │\n"
    "              │  streams text → s3://output-bucket    │\n"
    "              └──────────────────────────────────────┘"
)

doc.add_heading("Docker Services", level=3)
add_table(doc,
    ["Service", "Image", "Ports", "Purpose"],
    [
        ["localstack", "localstack/localstack:3.4", "4566", "SQS queues + S3 buckets"],
        ["rabbitmq", "rabbitmq:3.13-management-alpine", "5672, 15672", "RabbitMQ bus + Management UI"],
        ["zookeeper", "confluentinc/cp-zookeeper:7.6.1", "2181", "Kafka coordination"],
        ["kafka", "confluentinc/cp-kafka:7.6.1", "9092", "Kafka message bus"],
        ["ftp", "fauria/vsftpd", "21, 21100-21110", "Demo FTP server"],
        ["docconv-app", "Built from Dockerfile", "8080", "The conversion service"],
    ],
)

doc.add_heading("Project Structure", level=3)
add_code_block(doc,
    "docconv-service/\n"
    "├── docker-compose.yml        # Full infrastructure stack\n"
    "├── Dockerfile                # App container image\n"
    "├── requirements.txt          # Python dependencies\n"
    "├── .env.example              # All configurable env vars\n"
    "├── config/settings.py        # Pydantic settings (single source of truth)\n"
    "├── app/\n"
    "│   ├── main.py               # ★ Entry point – enables/disables components\n"
    "│   ├── models.py             # Pydantic models (ConversionJob, Result)\n"
    "│   ├── api.py                # FastAPI REST endpoints + test console\n"
    "│   ├── processor.py          # Pipeline: fetch → convert → upload\n"
    "│   ├── storage.py            # S3 output uploader\n"
    "│   ├── static/index.html     # API Test Console (HTML UI)\n"
    "│   ├── bus/\n"
    "│   │   ├── sqs_listener.py\n"
    "│   │   ├── rabbitmq_listener.py\n"
    "│   │   └── kafka_listener.py\n"
    "│   ├── fetchers/\n"
    "│   │   ├── dispatch.py       # Routes by location_type\n"
    "│   │   ├── s3_fetcher.py\n"
    "│   │   ├── url_fetcher.py\n"
    "│   │   └── ftp_fetcher.py\n"
    "│   └── converters/\n"
    "│       ├── dispatch.py       # Routes by document_type\n"
    "│       ├── pdf_converter.py  # PDF text + embedded image OCR\n"
    "│       ├── docx_converter.py\n"
    "│       ├── xlsx_converter.py # Also handles CSV\n"
    "│       ├── pptx_converter.py\n"
    "│       ├── html_converter.py\n"
    "│       ├── rtf_converter.py\n"
    "│       ├── odt_converter.py\n"
    "│       ├── text_converter.py\n"
    "│       └── image_converter.py\n"
    "├── scripts/\n"
    "│   ├── init-localstack.sh    # Creates SQS queue + S3 buckets\n"
    "│   └── demo.py               # Send test jobs via every channel\n"
    "└── tests/\n"
    "    └── test_full_flow.py     # Integration test suite (41 tests)"
)

# ═══════════════════════════════════════════════════════════
# 3. PREREQUISITES
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("3. Prerequisites", level=1)

add_table(doc,
    ["Requirement", "Version", "Notes"],
    [
        ["Docker", "20.10+", "Docker Desktop or Docker Engine"],
        ["Docker Compose", "v2.0+", "Included with Docker Desktop"],
        ["Disk space", "~3 GB", "For container images"],
        ["RAM", "4 GB minimum", "8 GB recommended for Kafka + LocalStack"],
        ["Python", "3.10+ (optional)", "Only needed for running demo.py or tests locally"],
    ],
)
add_note(doc, "All service dependencies (Tesseract, Poppler, etc.) are installed inside the Docker image.  You do not need to install them on your host.")

# ═══════════════════════════════════════════════════════════
# 4. QUICK START
# ═══════════════════════════════════════════════════════════

doc.add_heading("4. Quick Start – Docker Compose", level=1)

doc.add_heading("Step 1: Extract and enter the project", level=3)
add_code_block(doc, "unzip docconv-service.zip\ncd docconv-service")

doc.add_heading("Step 2: Make the init script executable", level=3)
add_code_block(doc, "chmod +x scripts/init-localstack.sh")

doc.add_heading("Step 3: Build and start all services", level=3)
add_code_block(doc, "docker compose up --build -d")
doc.add_paragraph(
    "This builds the app image, pulls infrastructure images, starts all 6 containers, "
    "and initializes LocalStack with the SQS queue and S3 buckets."
)

doc.add_heading("Step 4: Verify services are healthy", level=3)
add_code_block(doc, "docker compose ps")
doc.add_paragraph("All services should show \"healthy\" or \"running\" status.")

doc.add_heading("Step 5: Open the API Test Console", level=3)
doc.add_paragraph("Open your browser to http://localhost:8080 to access the built-in API test console.")

doc.add_heading("Step 6: Watch logs", level=3)
add_code_block(doc, "docker compose logs -f docconv-app")

doc.add_heading("Step 7: Run the demo script (optional)", level=3)
add_code_block(doc, "pip install boto3 pika confluent-kafka requests\npython scripts/demo.py")

doc.add_heading("Step 8: Stop the stack", level=3)
add_code_block(doc, "docker compose down          # Stop containers\ndocker compose down -v       # Stop and remove volumes")

# ═══════════════════════════════════════════════════════════
# 5. CONFIGURATION REFERENCE
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("5. Configuration Reference", level=1)
doc.add_paragraph(
    "All configuration is via environment variables.  Every variable has a sensible default. "
    "Copy .env.example to .env for local overrides."
)

doc.add_heading("Feature Flags", level=3)
add_table(doc,
    ["Variable", "Default", "Description"],
    [
        ["ENABLE_API", "true", "Start the FastAPI HTTP server on port 8080"],
        ["ENABLE_SQS", "true", "Start the SQS long-poll listener thread"],
        ["ENABLE_RABBITMQ", "true", "Start the RabbitMQ consumer thread"],
        ["ENABLE_KAFKA", "true", "Start the Kafka consumer thread"],
    ],
)
doc.add_paragraph('Set any flag to "false" to completely disable that component at startup.')

doc.add_heading("SQS Settings", level=3)
add_table(doc,
    ["Variable", "Default", "Description"],
    [
        ["SQS_ENDPOINT_URL", "http://localstack:4566", "SQS endpoint (LocalStack or real AWS)"],
        ["SQS_QUEUE_NAME", "docconv-jobs", "Name of the SQS queue to poll"],
        ["SQS_POLL_INTERVAL", "5", "Long-poll wait time in seconds"],
    ],
)

doc.add_heading("S3 Output Settings", level=3)
add_table(doc,
    ["Variable", "Default", "Description"],
    [
        ["S3_ENDPOINT_URL", "http://localstack:4566", "S3 endpoint for output uploads"],
        ["S3_OUTPUT_BUCKET", "docconv-output", "Default bucket for converted text"],
    ],
)

doc.add_heading("RabbitMQ Settings", level=3)
add_table(doc,
    ["Variable", "Default", "Description"],
    [
        ["RABBITMQ_HOST", "rabbitmq", "RabbitMQ server hostname"],
        ["RABBITMQ_PORT", "5672", "AMQP port"],
        ["RABBITMQ_USER", "docconv", "Authentication username"],
        ["RABBITMQ_PASS", "docconv", "Authentication password"],
        ["RABBITMQ_QUEUE", "docconv-jobs", "Queue name to consume"],
    ],
)

doc.add_heading("Kafka Settings", level=3)
add_table(doc,
    ["Variable", "Default", "Description"],
    [
        ["KAFKA_BOOTSTRAP_SERVERS", "kafka:29092", "Kafka broker address"],
        ["KAFKA_TOPIC", "docconv-jobs", "Topic to subscribe to"],
        ["KAFKA_GROUP_ID", "docconv-group", "Consumer group ID"],
    ],
)

doc.add_heading("AWS Credentials", level=3)
add_table(doc,
    ["Variable", "Default", "Description"],
    [
        ["AWS_ACCESS_KEY_ID", "test", "Used by boto3 for SQS and S3"],
        ["AWS_SECRET_ACCESS_KEY", "test", "Used by boto3 for SQS and S3"],
        ["AWS_DEFAULT_REGION", "us-east-1", "AWS region"],
    ],
)

doc.add_heading("Tuning", level=3)
add_table(doc,
    ["Variable", "Default", "Description"],
    [
        ["CHUNK_SIZE", "65536", "Read/write chunk size in bytes (64 KB)"],
        ["TMP_DIR", "/tmp/docconv", "Directory for temporary files"],
        ["LOG_LEVEL", "INFO", "Python logging level (DEBUG, INFO, WARNING, ERROR)"],
        ["API_PORT", "8080", "Port for the FastAPI HTTP server"],
    ],
)

# ═══════════════════════════════════════════════════════════
# 6. MESSAGE SCHEMA
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("6. Message Schema", level=1)
doc.add_paragraph(
    "Every message bus and the REST API accept the same JSON schema.  Fields are required "
    "or optional depending on the location_type."
)

doc.add_heading("Full JSON Schema", level=3)
add_code_block(doc,
    '{\n'
    '  "job_id": "optional-caller-id",\n'
    '  "document_type": "pdf",\n'
    '  "location_type": "s3",\n'
    '  \n'
    '  "s3_bucket": "my-bucket",\n'
    '  "s3_key": "path/to/file.pdf",\n'
    '  "s3_endpoint_url": null,\n'
    '  \n'
    '  "url": "https://example.com/doc.pdf",\n'
    '  \n'
    '  "ftp_host": "ftp.example.com",\n'
    '  "ftp_port": 21,\n'
    '  "ftp_path": "/path/to/file.pdf",\n'
    '  "ftp_user": "username",\n'
    '  "ftp_pass": "password",\n'
    '  \n'
    '  "auth_type": "none",\n'
    '  "auth_username": null,\n'
    '  "auth_password": null,\n'
    '  "auth_token": null,\n'
    '  \n'
    '  "output_s3_bucket": null,\n'
    '  "output_s3_key": null\n'
    '}'
)

doc.add_heading("Required Fields by Location Type", level=3)
add_table(doc,
    ["Field", "S3", "URL", "FTP", "Local (API)"],
    [
        ["document_type", "Required", "Required", "Required", "Required"],
        ["location_type", "Required", "Required", "Required", "Set automatically"],
        ["s3_bucket", "Required", "—", "—", "—"],
        ["s3_key", "Required", "—", "—", "—"],
        ["url", "—", "Required", "—", "—"],
        ["ftp_host", "—", "—", "Required", "—"],
        ["ftp_path", "—", "—", "Required", "—"],
    ],
)

doc.add_heading("Supported document_type values", level=3)
add_table(doc,
    ["Value", "Description", "Converter Module"],
    [
        ["pdf", "PDF documents (with embedded image OCR)", "pdf_converter.py"],
        ["docx", "Microsoft Word documents", "docx_converter.py"],
        ["xlsx", "Microsoft Excel spreadsheets", "xlsx_converter.py"],
        ["pptx", "Microsoft PowerPoint presentations", "pptx_converter.py"],
        ["html", "HTML web pages", "html_converter.py"],
        ["rtf", "Rich Text Format", "rtf_converter.py"],
        ["odt", "OpenDocument Text (LibreOffice)", "odt_converter.py"],
        ["txt", "Plain text (with encoding detection)", "text_converter.py"],
        ["csv", "Comma-separated values", "xlsx_converter.py"],
        ["image", "Images – JPEG, PNG, TIFF, BMP, WEBP (OCR)", "image_converter.py"],
    ],
)

# ═══════════════════════════════════════════════════════════
# 7. REST API REFERENCE
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("7. REST API Reference", level=1)
doc.add_paragraph("Base URL: http://localhost:8080  (configurable via API_PORT)")

doc.add_heading("GET /", level=3)
doc.add_paragraph("Returns the API Test Console HTML page.  Open in a browser to use the interactive UI.")

doc.add_heading("GET /health", level=3)
doc.add_paragraph("Returns the service status and enabled components.")
add_code_block(doc, 'curl http://localhost:8080/health')
doc.add_paragraph("Response:")
add_code_block(doc,
    '{\n'
    '  "status": "ok",\n'
    '  "api_enabled": true,\n'
    '  "sqs_enabled": true,\n'
    '  "rabbitmq_enabled": true,\n'
    '  "kafka_enabled": true\n'
    '}'
)

doc.add_heading("POST /convert/job", level=3)
doc.add_paragraph("Submit a conversion job.  The service fetches the document from the remote source, converts it, and uploads the text to S3.")
add_code_block(doc,
    'curl -X POST http://localhost:8080/convert/job \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{\n'
    '    "document_type": "html",\n'
    '    "location_type": "url",\n'
    '    "url": "https://example.com"\n'
    '  }\''
)
doc.add_paragraph("Response:")
add_code_block(doc,
    '{\n'
    '  "job_id": "a1b2c3d4e5f6",\n'
    '  "success": true,\n'
    '  "output_bucket": "docconv-output",\n'
    '  "output_key": "converted/a1b2c3d4e5f6.txt",\n'
    '  "error": null,\n'
    '  "characters_extracted": 1256\n'
    '}'
)

doc.add_heading("POST /convert/upload", level=3)
doc.add_paragraph("Upload a file directly for conversion (multipart form data).")
add_code_block(doc,
    'curl -X POST http://localhost:8080/convert/upload \\\n'
    '  -F "file=@report.pdf" \\\n'
    '  -F "document_type=pdf"'
)
doc.add_paragraph("Optional form fields: output_s3_bucket, output_s3_key.")

doc.add_heading("GET /docs", level=3)
doc.add_paragraph("FastAPI auto-generated Swagger/OpenAPI documentation.")

doc.add_heading("Disabling the API", level=3)
doc.add_paragraph("Set ENABLE_API=false to run the service without an HTTP endpoint.  The process will still consume from enabled message buses.")

# ═══════════════════════════════════════════════════════════
# 8. API TEST CONSOLE
# ═══════════════════════════════════════════════════════════

doc.add_heading("8. Using the API Test Console", level=1)
doc.add_paragraph(
    "The service includes a built-in web UI for testing all API endpoints.  "
    "Open http://localhost:8080 in your browser after starting the Docker stack."
)

doc.add_heading("Test Console Features", level=3)
add_bullet(doc, "Health Check tab – verify connectivity and see which components are enabled")
add_bullet(doc, "Upload tab – drag-and-drop or click to upload files for conversion, with auto-detection of document type from the file extension")
add_bullet(doc, "Job tab – build and send JSON conversion jobs with S3, URL, or FTP sources, with a live JSON preview panel")
add_bullet(doc, "Schema tab – reference for the full message schema and all supported values")
add_bullet(doc, "Request history – logs every request with method, path, status code, and response time")
add_bullet(doc, "Syntax-highlighted JSON responses with result summary cards")

doc.add_heading("Configuration", level=3)
doc.add_paragraph(
    "The Base URL field at the top defaults to http://localhost:8080.  Change it if "
    "your service is running on a different host or port.  Click 'Test Connection' to verify."
)

# ═══════════════════════════════════════════════════════════
# 9. MESSAGE BUS INTEGRATION
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("9. Message Bus Integration", level=1)
doc.add_paragraph(
    "The service can consume jobs from SQS, RabbitMQ, and Kafka simultaneously.  Each listener "
    "runs in its own daemon thread and can be independently enabled or disabled."
)

doc.add_heading("9.1  Amazon SQS", level=2)
doc.add_paragraph(
    "The SQS listener long-polls the configured queue.  In the Docker Compose stack, "
    "LocalStack provides a local SQS emulation.  For production, point SQS_ENDPOINT_URL "
    "to the real AWS endpoint (or remove it to use the default)."
)
doc.add_heading("Sending a test message via AWS CLI", level=3)
add_code_block(doc,
    'aws --endpoint-url=http://localhost:4566 sqs send-message \\\n'
    '  --queue-url http://localhost:4566/000000000000/docconv-jobs \\\n'
    '  --message-body \'{\n'
    '    "document_type": "txt",\n'
    '    "location_type": "s3",\n'
    '    "s3_bucket": "docconv-input",\n'
    '    "s3_key": "test.txt"\n'
    '  }\''
)

doc.add_heading("Sending via Python (boto3)", level=3)
add_code_block(doc,
    'import boto3, json\n'
    'sqs = boto3.client("sqs",\n'
    '    endpoint_url="http://localhost:4566",\n'
    '    aws_access_key_id="test",\n'
    '    aws_secret_access_key="test",\n'
    '    region_name="us-east-1")\n'
    'queue_url = sqs.get_queue_url(QueueName="docconv-jobs")["QueueUrl"]\n'
    'sqs.send_message(\n'
    '    QueueUrl=queue_url,\n'
    '    MessageBody=json.dumps({\n'
    '        "document_type": "html",\n'
    '        "location_type": "url",\n'
    '        "url": "https://example.com"\n'
    '    }))'
)

doc.add_heading("9.2  RabbitMQ", level=2)
doc.add_paragraph(
    "The RabbitMQ listener connects via AMQP and consumes from the configured queue. "
    "The management UI is available at http://localhost:15672 (user: docconv, pass: docconv)."
)
doc.add_heading("Sending via Python (pika)", level=3)
add_code_block(doc,
    'import pika, json\n'
    'creds = pika.PlainCredentials("docconv", "docconv")\n'
    'conn = pika.BlockingConnection(\n'
    '    pika.ConnectionParameters("localhost", 5672, credentials=creds))\n'
    'ch = conn.channel()\n'
    'ch.queue_declare(queue="docconv-jobs", durable=True)\n'
    'ch.basic_publish(\n'
    '    exchange="",\n'
    '    routing_key="docconv-jobs",\n'
    '    body=json.dumps({"document_type":"pdf","location_type":"url",\n'
    '                     "url":"https://example.com/doc.pdf"}),\n'
    '    properties=pika.BasicProperties(delivery_mode=2))\n'
    'conn.close()'
)

doc.add_heading("9.3  Apache Kafka", level=2)
doc.add_paragraph(
    "The Kafka listener subscribes to the configured topic using the confluent-kafka consumer. "
    "Auto-topic creation is enabled in the Docker Compose stack."
)
doc.add_heading("Sending via Python (confluent-kafka)", level=3)
add_code_block(doc,
    'from confluent_kafka import Producer\n'
    'import json\n'
    'producer = Producer({"bootstrap.servers": "localhost:9092"})\n'
    'producer.produce("docconv-jobs",\n'
    '    value=json.dumps({\n'
    '        "document_type": "docx",\n'
    '        "location_type": "url",\n'
    '        "url": "https://example.com/report.docx"\n'
    '    }).encode("utf-8"))\n'
    'producer.flush()'
)

doc.add_heading("Sending via Kafka CLI", level=3)
add_code_block(doc,
    'echo \'{"document_type":"txt","location_type":"url","url":"https://example.com"}\' | \\\n'
    '  docker exec -i docconv-kafka kafka-console-producer \\\n'
    '    --bootstrap-server localhost:29092 \\\n'
    '    --topic docconv-jobs'
)

# ═══════════════════════════════════════════════════════════
# 10. DOCUMENT CONVERTERS
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("10. Document Converters", level=1)
doc.add_paragraph(
    "Each converter is a standalone Python module exposing a convert_to_text(file_path) "
    "function that returns a generator of text chunks.  This design keeps memory usage "
    "constant regardless of document size."
)

converters = [
    ("PDF (pdf_converter.py)", 
     "Extracts selectable text via pdfplumber page-by-page. If a page has no selectable "
     "text, it falls back to full-page OCR by rendering the page as an image with Poppler's "
     "pdf2image and running Tesseract. Additionally, embedded images on every page are "
     "extracted and OCR'd separately, with results tagged as [IMAGE TEXT]."),
    ("DOCX (docx_converter.py)",
     "Reads paragraphs in batches of 50 using python-docx.  Also extracts all table content "
     "with cell delimiters.  Yields text in chunks."),
    ("XLSX / CSV (xlsx_converter.py)",
     "For XLSX: reads each sheet with openpyxl in read-only mode, yielding rows in batches "
     "of 200 with pipe-delimited cells.  For CSV: uses the csv stdlib module with the same "
     "batching strategy."),
    ("PPTX (pptx_converter.py)",
     "Reads each slide using python-pptx.  Extracts text from all shapes (text frames "
     "and tables) plus speaker notes.  Yields one slide at a time."),
    ("HTML (html_converter.py)",
     "Strips script/style/noscript tags using BeautifulSoup + lxml.  Extracts visible text "
     "with newline separators.  Yields in configurable chunk sizes."),
    ("RTF (rtf_converter.py)",
     "Strips RTF control words using the striprtf library.  Yields plain text in chunks."),
    ("ODT (odt_converter.py)",
     "Reads OpenDocument Text paragraphs via odfpy.  Yields in batches of 50 paragraphs."),
    ("TXT (text_converter.py)",
     "Detects encoding via chardet, then reads and yields text in CHUNK_SIZE chunks. "
     "Handles UTF-8, Latin-1, Windows-1252, and other encodings."),
]

for title, desc in converters:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

# ═══════════════════════════════════════════════════════════
# 11. IMAGE OCR
# ═══════════════════════════════════════════════════════════

doc.add_heading("11. Image OCR", level=1)
doc.add_paragraph(
    "The image converter (image_converter.py) handles JPEG, PNG, TIFF (including multi-frame), "
    "BMP, and WEBP formats.  It uses Pillow to open the image and Tesseract for OCR."
)

doc.add_heading("Processing Steps", level=3)
add_bullet(doc, "Open the image with Pillow")
add_bullet(doc, "For multi-frame TIFFs, iterate through each frame")
add_bullet(doc, "Down-scale if either dimension exceeds 4000px (prevents OOM)")
add_bullet(doc, "Convert to RGB if needed (handles RGBA, palette modes)")
add_bullet(doc, "Run Tesseract OCR and yield the extracted text")

doc.add_heading("PDF Image Extraction", level=3)
doc.add_paragraph(
    "The PDF converter additionally extracts images embedded in each page using "
    "pdfplumber's image detection.  Each image is cropped from the page, rendered "
    "at 200 DPI, and OCR'd.  Results are tagged with [IMAGE TEXT] prefix."
)

# ═══════════════════════════════════════════════════════════
# 12. FILE FETCHERS
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("12. File Fetchers (S3, URL, FTP)", level=1)
doc.add_paragraph(
    "Each fetcher downloads a document to a local temporary file using chunked streaming. "
    "All fetchers use Python generators with yield to ensure the tmp file is cleaned up "
    "after processing, even if an error occurs."
)

doc.add_heading("S3 Fetcher (s3_fetcher.py)", level=3)
add_bullet(doc, "Uses boto3 to get the S3 object")
add_bullet(doc, "Reads the response body in CHUNK_SIZE chunks")
add_bullet(doc, "Writes to a NamedTemporaryFile in TMP_DIR")
add_bullet(doc, "Supports custom endpoint URLs for LocalStack/MinIO")

doc.add_heading("URL Fetcher (url_fetcher.py)", level=3)
add_bullet(doc, "Uses requests with stream=True for chunked download")
add_bullet(doc, "Supports auth_type=none, basic (username/password), and bearer (token)")
add_bullet(doc, "Timeout set to 120 seconds")

doc.add_heading("FTP Fetcher (ftp_fetcher.py)", level=3)
add_bullet(doc, "Uses stdlib ftplib with retrbinary for binary download")
add_bullet(doc, "Supports anonymous and authenticated access")
add_bullet(doc, "Reads in CHUNK_SIZE blocks")

# ═══════════════════════════════════════════════════════════
# 13. OUTPUT STORAGE
# ═══════════════════════════════════════════════════════════

doc.add_heading("13. Output Storage (S3)", level=1)
doc.add_paragraph(
    "Converted text is uploaded to the S3 output bucket.  The upload process writes "
    "yielded text chunks to a local tmp file first, then uploads the complete file to S3. "
    "This approach avoids buffering large text strings in memory and handles S3 upload "
    "retries gracefully."
)

doc.add_heading("Default Output Path", level=3)
add_code_block(doc, "s3://{S3_OUTPUT_BUCKET}/converted/{job_id}.txt")
doc.add_paragraph("Override with output_s3_bucket and output_s3_key in the job message.")

doc.add_heading("Checking Output", level=3)
add_code_block(doc,
    '# List converted files\n'
    'aws --endpoint-url=http://localhost:4566 s3 ls \\\n'
    '  s3://docconv-output/converted/ --recursive\n\n'
    '# View a converted file\n'
    'aws --endpoint-url=http://localhost:4566 s3 cp \\\n'
    '  s3://docconv-output/converted/a1b2c3d4e5f6.txt -'
)

# ═══════════════════════════════════════════════════════════
# 14. MEMORY EFFICIENCY
# ═══════════════════════════════════════════════════════════

doc.add_heading("14. Memory Efficiency & Streaming", level=1)
doc.add_paragraph("The service is designed for low memory usage when processing large documents.")

add_bullet(doc, "Fetchers: ", "Fetchers: ")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Stream remote files to disk in CHUNK_SIZE chunks using Python generators (yield).  The file is written to TMP_DIR and auto-deleted after processing.")

add_bullet(doc, "Converters: ", "Converters: ")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Yield text one page/chunk at a time.  They never build the full text string in memory.  PDF pages are processed and discarded individually.")

add_bullet(doc, "Upload: ", "Upload: ")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Writes yielded chunks to a tmp file, then uploads that file to S3.  No full-text buffering in RAM.")

add_bullet(doc, "PDF images: ", "PDF images: ")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Rendered to tmp PNG files for OCR, deleted immediately after each page is processed.")

add_bullet(doc, "API uploads: ", "API uploads: ")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Streamed to disk in chunks from the HTTP request body, never held in memory.")

# ═══════════════════════════════════════════════════════════
# 15. TESTING
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("15. Testing", level=1)

doc.add_heading("Running the Integration Test Suite", level=3)
add_code_block(doc, "cd docconv-service\npip install -r requirements.txt reportlab httpx\npython -m tests.test_full_flow")

doc.add_heading("What the Tests Cover (41 tests)", level=3)
add_table(doc,
    ["Section", "Count", "Description"],
    [
        ["Sample Generation", "10", "Creates sample files for every supported format"],
        ["Converter Unit Tests", "10", "Tests each converter module independently"],
        ["Dispatcher Tests", "10", "Verifies DocumentType → converter routing"],
        ["Pipeline Tests", "10", "End-to-end: file → convert → output"],
        ["Model Validation", "4", "Pydantic schema validation and rejection"],
        ["REST API Tests", "7", "Health, upload (5 formats), validation"],
    ],
)

doc.add_heading("Running via Docker", level=3)
add_code_block(doc,
    'docker compose exec docconv-app python -m tests.test_full_flow'
)

# ═══════════════════════════════════════════════════════════
# 16. TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════

doc.add_heading("16. Troubleshooting", level=1)

problems = [
    ("Services fail to start",
     'Run "docker compose logs <service>" to check for errors.  Ensure ports 4566, 5672, '
     '9092, 8080 are not in use.  Increase Docker memory to 8 GB if Kafka/Zookeeper OOM.'),
    ("SQS listener says 'queue not found'",
     "LocalStack may still be initializing.  The listener retries 30 times with 2-second "
     "intervals.  Check that init-localstack.sh is executable (chmod +x)."),
    ("OCR returns empty text",
     "Ensure the tesseract-ocr and tesseract-ocr-eng packages are installed in the Docker "
     "image (they are by default).  Check the image quality – very low resolution or "
     "handwritten text may produce poor results."),
    ("PDF conversion is slow",
     "Full-page OCR (for scanned PDFs) renders each page at 200 DPI.  This is CPU-intensive. "
     "For faster processing of selectable-text PDFs, no OCR is triggered."),
    ("Upload fails with 'Connection refused'",
     "The S3 upload is targeting LocalStack.  Ensure LocalStack is running and healthy.  "
     'Check with: curl http://localhost:4566/_localstack/health'),
    ("RabbitMQ connection refused",
     "The listener auto-retries every 5 seconds.  RabbitMQ may take 10-15 seconds to start.  "
     "Check: docker compose logs rabbitmq"),
    ("Kafka consumer not receiving messages",
     "Ensure the topic exists or auto-creation is enabled (KAFKA_AUTO_CREATE_TOPICS_ENABLE=true).  "
     "Check bootstrap server address matches between producer and consumer."),
    ("API Test Console shows 'Disconnected'",
     "Click 'Test Connection' to retry.  Ensure the Base URL matches your Docker host. "
     "If running Docker in a VM, use the VM's IP instead of localhost."),
]

for title, solution in problems:
    doc.add_heading(title, level=3)
    doc.add_paragraph(solution)

# ═══════════════════════════════════════════════════════════
# 17. CURL EXAMPLES
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("17. curl Examples Cookbook", level=1)

examples = [
    ("Health check",
     'curl http://localhost:8080/health'),
    ("Upload a PDF",
     'curl -X POST http://localhost:8080/convert/upload \\\n'
     '  -F "file=@document.pdf" \\\n'
     '  -F "document_type=pdf"'),
    ("Upload an image for OCR",
     'curl -X POST http://localhost:8080/convert/upload \\\n'
     '  -F "file=@scan.png" \\\n'
     '  -F "document_type=image"'),
    ("Upload a DOCX with custom output key",
     'curl -X POST http://localhost:8080/convert/upload \\\n'
     '  -F "file=@report.docx" \\\n'
     '  -F "document_type=docx" \\\n'
     '  -F "output_s3_key=reports/2025/q1.txt"'),
    ("Convert from URL",
     'curl -X POST http://localhost:8080/convert/job \\\n'
     '  -H "Content-Type: application/json" \\\n'
     '  -d \'{"document_type":"html","location_type":"url",\n'
     '       "url":"https://example.com"}\''),
    ("Convert from URL with Basic auth",
     'curl -X POST http://localhost:8080/convert/job \\\n'
     '  -H "Content-Type: application/json" \\\n'
     '  -d \'{"document_type":"pdf","location_type":"url",\n'
     '       "url":"https://secure.example.com/doc.pdf",\n'
     '       "auth_type":"basic",\n'
     '       "auth_username":"user","auth_password":"pass"}\''),
    ("Convert from URL with Bearer token",
     'curl -X POST http://localhost:8080/convert/job \\\n'
     '  -H "Content-Type: application/json" \\\n'
     '  -d \'{"document_type":"xlsx","location_type":"url",\n'
     '       "url":"https://api.example.com/export.xlsx",\n'
     '       "auth_type":"bearer","auth_token":"tok_abc123"}\''),
    ("Convert from S3 (LocalStack)",
     'curl -X POST http://localhost:8080/convert/job \\\n'
     '  -H "Content-Type: application/json" \\\n'
     '  -d \'{"document_type":"pdf","location_type":"s3",\n'
     '       "s3_bucket":"docconv-input","s3_key":"docs/report.pdf",\n'
     '       "s3_endpoint_url":"http://localstack:4566"}\''),
    ("Convert from FTP",
     'curl -X POST http://localhost:8080/convert/job \\\n'
     '  -H "Content-Type: application/json" \\\n'
     '  -d \'{"document_type":"txt","location_type":"ftp",\n'
     '       "ftp_host":"ftp","ftp_path":"/data/file.txt",\n'
     '       "ftp_user":"docconv","ftp_pass":"docconv"}\''),
    ("List converted output files",
     'aws --endpoint-url=http://localhost:4566 s3 ls \\\n'
     '  s3://docconv-output/converted/ --recursive'),
    ("Download and view converted text",
     'aws --endpoint-url=http://localhost:4566 s3 cp \\\n'
     '  s3://docconv-output/converted/JOB_ID.txt -'),
]

for title, cmd in examples:
    doc.add_heading(title, level=3)
    add_code_block(doc, cmd)

# ── Save ─────────────────────────────────────────────────
output_path = "/home/claude/docconv-service/DocConv_User_Guide.docx"
doc.save(output_path)
print(f"User guide saved to {output_path}")
